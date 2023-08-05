# -*- coding: utf-8 -*-

from __future__ import unicode_literals

import os
import sys

sys.path.append(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
)

from docx import Document
from docx.oxml.table import CT_Row, CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.parfmt import CT_PPr, CT_TabStop
from docx.table import Table
from jiexiproject.util.const_def import nsmap


def get_p_style(path):
    doc = Document(path)
    root = doc._body
    # 获取p的格式信息
    a = [nsmap["w:hangingChars"], nsmap["w:firstLineChars"], nsmap["w:leftChars"]]
    for i in root.paragraphs:
        p = i._element.pPr.ind
        for j in a:
            if p.get(j) and p.get(j) != "0":
                if j.split("}")[1] == "leftChars":
                    for node in i.runs:
                        for kk in node._element.t_lst:
                            if kk.get(nsmap["w:t"]):
                                print(kk.get(nsmap["w:t"]))
                                print(node.text)


def Get_Size_Of_File(filepath):
    f = open(filepath, mode="rb")
    f.seek(0, 2)
    filesize = f.tell()
    print(filesize)
    f.close()


def get_choice(path):
    choice = ("A", "B", "C", "D")
    doc = Document(path)
    root = doc._body
    # 获取p的格式信息
    a = [nsmap["w:hangingChars"], nsmap["w:firstLineChars"], nsmap["w:leftChars"]]
    for i in root.paragraphs:
        p = i._element.pPr.ind
        for j in a:
            res = []
            ress = []
            if p.get(j) and p.get(j) != "0":
                if (
                    j.split("}")[1] == "firstLineChars"
                    and len(i._element.pPr.tabs.tab_lst) == 1
                ):
                    for node in i.runs:
                        if len(node._element) == 1 and isinstance(
                            node._element[0], CT_TabStop
                        ):
                            if ress:
                                res.append(ress)
                                ress = []
                        else:
                            ress.append(node)
                    if ress:
                        res.append(ress)
                    for kk in res:
                        for kkk in kk:
                            if kkk.text.startswith(choice):
                                print(kkk.text)
                                print("*" * 20)


def get_xml_body(path):
    from zipfile import ZipFile

    from lxml import etree

    test = ZipFile(path).read("word/document.xml")
    doc = etree.fromstring(test)
    with open(r"E:/一起docx/jiexi_docx/test.xml", mode="wb") as f:
        f.write(
            etree.tostring(doc, encoding="utf-8", with_tail=True, pretty_print=True)
        )


def get_underline(path):
    doc = Document(path)
    root = doc._body
    for p in root.paragraphs:
        for r in p.runs:
            if r.underline:
                print(list(r._element.rPr))
                # if not r.text == "       ":
                #     print(r.text)


def get_b(path):
    doc = Document(path)
    root = doc._body.paragraphs
    for p in root:
        for r in p.runs:
            if r.bold:
                print(r.text)


def get_dot(path):
    doc = Document(path)
    root = doc._body.paragraphs
    for p in root:
        for r in p.runs:
            for i in r._element.iter():
                if i.tag == nsmap.get("w:em"):
                    if i.get(nsmap["w:val"]) == "dot":
                        print("ok")


def get_cell(path):
    doc = Document(path)
    root = doc._body._element
    res = []
    for node in root:
        if isinstance(node, CT_Tbl):
            for row in node:
                if isinstance(row, CT_Row):
                    for col in row:
                        res.append(list(col[0])[1:])
            print("*" * 20)
    print(res)


def get_cell(path):
    doc = Document(path)
    root = doc._body
    res = []
    for node in root.tables:
        if isinstance(node, Table):
            for row in node.rows:
                for i in row.cells[0]._element:
                    if isinstance(i, CT_P):
                        print(i)
                print("*" * 20)


def get_cell(path):
    doc = Document(path)
    root = doc._body
    res = []
    for node in root.tables:
        for row in node.rows:
            for cell in row.cells:
                print(cell.paragraphs[0].text)


if __name__ == "__main__":
    # path = r"D:\解析文档\docx_jiexi\数学、英语docx\初中语文箐优网\RES_000203396287初中语文.docx"
    # get_p_style(path)
    path = r"D:/解析文档/docx_jiexi/数学、英语docx/小学英语/RES_000232261165.docx"
    get_xml_body(path)
